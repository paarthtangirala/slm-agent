"""
Advanced document processing for SLM Personal Agent
Supports PDF, DOCX, XLSX, images, and more
"""

import os
import io
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
import asyncio
from datetime import datetime
import hashlib
import mimetypes

# Document processing imports
import PyPDF2
from docx import Document
import openpyxl
from PIL import Image
import pytesseract

# FastAPI and async imports
import aiofiles
from fastapi import UploadFile

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Supported file types
        self.supported_types = {
            'text': ['.txt', '.md', '.py', '.js', '.json', '.yaml', '.yml', '.css', '.html', '.xml'],
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'excel': ['.xlsx', '.xls'], 
            'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'],
            'code': ['.py', '.js', '.tsx', '.jsx', '.ts', '.go', '.rs', '.cpp', '.c', '.java']
        }
    
    def get_file_type(self, filename: str) -> str:
        """Determine file type based on extension"""
        ext = Path(filename).suffix.lower()
        
        for file_type, extensions in self.supported_types.items():
            if ext in extensions:
                return file_type
        return 'unknown'
    
    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported"""
        return self.get_file_type(filename) != 'unknown'
    
    async def save_upload(self, file: UploadFile) -> Dict[str, str]:
        """Save uploaded file and return metadata"""
        try:
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = "".join(c for c in file.filename if c.isalnum() or c in '._-')
            unique_filename = f"{timestamp}_{safe_filename}"
            file_path = self.upload_dir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)
            
            # Generate file hash for deduplication
            file_hash = hashlib.md5(content).hexdigest()
            
            return {
                "filename": file.filename,
                "saved_as": unique_filename,
                "path": str(file_path),
                "size": len(content),
                "type": self.get_file_type(file.filename),
                "hash": file_hash,
                "uploaded_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error saving file {file.filename}: {e}")
            raise
    
    async def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type == 'text' or file_type == 'code':
                return await self._extract_text_file(file_path)
            elif file_type == 'pdf':
                return await self._extract_pdf(file_path)
            elif file_type == 'word':
                return await self._extract_docx(file_path)
            elif file_type == 'excel':
                return await self._extract_xlsx(file_path)
            elif file_type == 'image':
                return await self._extract_image_ocr(file_path)
            else:
                return f"Unsupported file type: {file_type}"
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"Error processing file: {str(e)}"
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return await f.read()
    
    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text_content = []
        
        def extract_pdf_sync():
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                    except Exception as e:
                        text_content.append(f"--- Page {page_num + 1} (Error) ---\nError extracting page: {e}")
        
        # Run in thread to avoid blocking
        await asyncio.get_event_loop().run_in_executor(None, extract_pdf_sync)
        
        return "\n\n".join(text_content) if text_content else "No text content found in PDF"
    
    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word documents"""
        def extract_docx_sync():
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_content.append("--- Table ---\n" + "\n".join(table_text))
            
            return "\n\n".join(text_content)
        
        return await asyncio.get_event_loop().run_in_executor(None, extract_docx_sync)
    
    async def _extract_xlsx(self, file_path: str) -> str:
        """Extract text from Excel files"""
        def extract_xlsx_sync():
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = [f"--- Sheet: {sheet_name} ---"]
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):  # Skip empty rows
                        sheet_content.append(" | ".join(row_data))
                
                if len(sheet_content) > 1:  # More than just the header
                    text_content.extend(sheet_content)
            
            return "\n".join(text_content)
        
        return await asyncio.get_event_loop().run_in_executor(None, extract_xlsx_sync)
    
    async def _extract_image_ocr(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        def ocr_sync():
            try:
                # Open and process image
                with Image.open(file_path) as img:
                    # Convert to RGB if necessary
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    # Perform OCR
                    text = pytesseract.image_to_string(img, lang='eng')
                    return text.strip() if text.strip() else "No text detected in image"
                    
            except Exception as e:
                return f"OCR failed: {str(e)}. Note: Make sure Tesseract is installed."
        
        return await asyncio.get_event_loop().run_in_executor(None, ocr_sync)
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks for better embedding"""
        if not text.strip():
            return []
        
        # Split by sentences first, then by words if needed
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip() + '. ' if not sentence.endswith('.') else sentence.strip() + ' '
            
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # If no sentence-based chunking worked, fall back to word chunking
        if not chunks and text:
            words = text.split()
            for i in range(0, len(words), chunk_size - overlap):
                chunk = " ".join(words[i:i + chunk_size])
                chunks.append(chunk)
        
        return chunks
    
    async def get_file_info(self, file_path: str) -> Dict:
        """Get detailed information about a file"""
        try:
            path = Path(file_path)
            stat = path.stat()
            
            return {
                "name": path.name,
                "size": stat.st_size,
                "type": self.get_file_type(path.name),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "path": str(path),
                "exists": path.exists()
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return {"error": str(e)}
    
    async def delete_file(self, file_path: str) -> bool:
        """Delete a file"""
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
    
    async def list_uploaded_files(self) -> List[Dict]:
        """List all uploaded files with metadata"""
        files = []
        try:
            for file_path in self.upload_dir.iterdir():
                if file_path.is_file():
                    info = await self.get_file_info(str(file_path))
                    files.append(info)
        except Exception as e:
            logger.error(f"Error listing files: {e}")
        
        return sorted(files, key=lambda x: x.get('modified', ''), reverse=True)

# Global document processor instance
doc_processor = DocumentProcessor()